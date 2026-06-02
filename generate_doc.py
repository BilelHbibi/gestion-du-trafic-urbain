from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page setup ──────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.27)   # A4
section.page_height = Inches(11.69)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── Styles helper ───────────────────────────────────────────────
def set_font(run, name='Calibri', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=(30, 100, 200))
    p.space_before = Pt(18)
    p.space_after  = Pt(6)
    # underline
    run.font.underline = True
    return p

def heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=(0, 70, 160))
    p.space_before = Pt(12)
    p.space_after  = Pt(4)
    return p

def heading3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=(50, 50, 50))
    p.space_before = Pt(8)
    p.space_after  = Pt(2)
    return p

def body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    p.space_after = Pt(2)
    # parse bold (**...**)
    import re
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            set_font(r, bold=True, size=11)
        else:
            r = p.add_run(part)
            set_font(r, size=11)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_font(run, bold=True, size=10, color=(255,255,255))
        # blue background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1E64C8')
        tcPr.append(shd)

    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                set_font(run, size=10)
            # alternate row color
            if ri % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'E8F0FE')
                tcPr.append(shd)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def separator():
    p = doc.add_paragraph()
    p.space_after = Pt(2)
    run = p.add_run('─' * 80)
    run.font.color.rgb = RGBColor(200, 200, 200)
    run.font.size = Pt(8)

# ════════════════════════════════════════════════════════════════
#  PAGE DE TITRE
# ════════════════════════════════════════════════════════════════
# Logo / titre principal
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(20)
r = p.add_run('🚦 TRAFFIC PLATFORM')
set_font(r, size=28, bold=True, color=(30, 100, 200))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Plateforme Intelligente de Gestion du Trafic Urbain')
set_font(r, size=16, bold=False, color=(80, 80, 80))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Mini Projet — Module Web Services & GraphQL')
set_font(r, size=12, bold=True, color=(50, 50, 50))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('TEK-UP University — 2024/2025')
set_font(r, size=12, color=(100, 100, 100))

doc.add_paragraph()
separator()
doc.add_paragraph()

# Infos
info_table = doc.add_table(rows=4, cols=2)
info_table.style = 'Table Grid'
infos = [
    ('Étudiant', 'Bilel Hbibi'),
    ('Module', 'Web Services & GraphQL'),
    ('Encadrant', 'Prof. TEK-UP'),
    ('GitHub', 'github.com/BilelHbibi/gestion-du-trafic-urbain'),
]
for i, (label, val) in enumerate(infos):
    row = info_table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = val
    for run in row.cells[0].paragraphs[0].runs:
        set_font(run, bold=True, size=11)
    for run in row.cells[1].paragraphs[0].runs:
        set_font(run, size=11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  TABLE DES MATIÈRES (manuelle)
# ════════════════════════════════════════════════════════════════
heading1('Table des matières')
toc_items = [
    '1. Description générale du projet',
    '2. Objectifs du projet',
    '3. Architecture de la plateforme',
    '4. Technologies utilisées',
    '5. Les 5 Services Microservices',
    '   5.1 Service Authentification',
    '   5.2 Service Gestion des Véhicules',
    '   5.3 Service Gestion du Trafic',
    '   5.4 Service Gestion des Incidents',
    '   5.5 Service Notifications',
    '6. API Gateway GraphQL',
    '7. WebSocket — Temps Réel',
    '8. Base de données',
    '9. Sécurité',
    '10. Docker & Déploiement',
    '11. Guide de test complet',
    '12. Résumé des fonctionnalités',
]
for item in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(item)
    set_font(r, size=11, color=(30, 100, 200) if not item.startswith('   ') else (80, 80, 80))
    p.space_after = Pt(3)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  1. DESCRIPTION GÉNÉRALE
# ════════════════════════════════════════════════════════════════
heading1('1. Description générale du projet')

body(
    'Traffic Platform est une application web distribuée développée dans le cadre du module '
    'Web Services. Elle simule une plateforme intelligente de gestion du trafic urbain, '
    'telle que celles utilisées par les municipalités et les autorités routières dans les '
    'grandes villes modernes.'
)
body(
    'Le projet repose sur une architecture microservices : chaque fonctionnalité est '
    'développée comme un service indépendant, avec sa propre base de données. Ces services '
    'communiquent entre eux via des requêtes HTTP, et sont tous accessibles depuis un point '
    'd\'entrée unique : l\'API Gateway GraphQL.'
)
body(
    'La plateforme permet de superviser les véhicules en circulation, détecter les zones '
    'congestionnées, signaler des incidents routiers, et envoyer des notifications en temps '
    'réel aux opérateurs connectés.'
)

separator()

# ════════════════════════════════════════════════════════════════
#  2. OBJECTIFS
# ════════════════════════════════════════════════════════════════
heading1('2. Objectifs du projet')

heading2('Objectifs techniques')
bullet('**Maîtriser l\'architecture microservices** : chaque service est indépendant et déployable séparément')
bullet('**Implémenter GraphQL** : une API flexible qui remplace les APIs REST classiques')
bullet('**Sécuriser l\'application** : authentification JWT, contrôle des rôles, protection contre les attaques')
bullet('**Communiquer en temps réel** : utilisation de WebSocket pour les alertes instantanées')
bullet('**Conteneuriser l\'application** : Docker Compose pour un déploiement en une seule commande')

heading2('Objectifs fonctionnels')
bullet('Permettre à des administrateurs et opérateurs de gérer le trafic urbain')
bullet('Suivre les véhicules avec leur position GPS simulée')
bullet('Détecter automatiquement les zones congestionnées selon la densité du trafic')
bullet('Signaler et suivre les incidents routiers (accidents, travaux, embouteillages)')
bullet('Notifier les utilisateurs en temps réel dès qu\'un incident est déclaré')

separator()

# ════════════════════════════════════════════════════════════════
#  3. ARCHITECTURE
# ════════════════════════════════════════════════════════════════
heading1('3. Architecture de la plateforme')

heading2('Vue d\'ensemble')
body(
    'L\'application suit une architecture microservices avec un API Gateway central. '
    'Le client (navigateur ou application) envoie ses requêtes uniquement vers le Gateway '
    'GraphQL (port 4000). Le Gateway se charge ensuite de router ces requêtes vers le bon '
    'service interne.'
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    '[ Client / Navigateur ]\n'
    '         ↓  requêtes GraphQL\n'
    '[ API Gateway GraphQL — port 4000 ]\n'
    '    ↓         ↓         ↓         ↓         ↓\n'
    '[ Auth ] [ Vehicle ] [ Traffic ] [ Incident ] [ Notif ]\n'
    ' :3001    :3002       :3003       :3004         :3005\n'
    '   ↓         ↓         ↓           ↓              ↓\n'
    'db_auth  db_vehicles db_traffic db_incidents db_notifications'
)
set_font(r, name='Courier New', size=9, color=(50, 50, 50))

heading2('Principe de l\'isolation')
body(
    'Chaque service possède sa propre base de données MySQL. Cela signifie qu\'ils ne '
    'partagent pas de données directement : toute communication passe par des appels HTTP '
    'entre services. C\'est le principe fondamental des microservices.'
)

heading2('Types de communication')
add_table(
    ['Type', 'Utilisation', 'Exemple'],
    [
        ['GraphQL (HTTP)', 'Client → Gateway', 'Requête pour voir les véhicules'],
        ['REST (HTTP interne)', 'Service → Service', 'Incident notifie Notification'],
        ['WebSocket', 'Gateway → Client', 'Alerte incident en temps réel'],
    ],
    col_widths=[4, 5, 7]
)

doc.add_paragraph()
separator()

# ════════════════════════════════════════════════════════════════
#  4. TECHNOLOGIES
# ════════════════════════════════════════════════════════════════
heading1('4. Technologies utilisées')

add_table(
    ['Catégorie', 'Technologie', 'Rôle dans le projet'],
    [
        ['Runtime', 'Node.js v18', 'Exécution JavaScript côté serveur'],
        ['Framework Web', 'Express.js', 'Création des APIs REST (services internes)'],
        ['GraphQL', 'Apollo Server v4', 'API Gateway avec schéma GraphQL complet'],
        ['Base de données', 'MySQL 8 (XAMPP)', 'Stockage relationnel, 5 bases isolées'],
        ['Authentification', 'JWT (jsonwebtoken)', 'Tokens sécurisés, expiration 24h'],
        ['Chiffrement', 'bcryptjs', 'Hash des mots de passe (sel + hachage)'],
        ['Communication', 'Axios', 'Appels HTTP entre services'],
        ['Temps réel', 'Socket.io', 'WebSocket pour alertes instantanées'],
        ['Frontend', 'React + Vite', 'Dashboard interactif'],
        ['Déploiement', 'Docker Compose', 'Conteneurisation de toute la plateforme'],
        ['Tests', 'Postman', 'Collection complète de tests API'],
    ],
    col_widths=[3.5, 4, 8.5]
)

doc.add_paragraph()
separator()

# ════════════════════════════════════════════════════════════════
#  5. LES 5 SERVICES
# ════════════════════════════════════════════════════════════════
heading1('5. Les 5 Services Microservices')

# 5.1 AUTH
heading2('5.1  Service Authentification  (port 3001)')
body(
    'Ce service gère tout ce qui concerne les utilisateurs : leur inscription, leur connexion, '
    'et la vérification de leur identité. C\'est le service le plus important car tous les '
    'autres services lui délèguent la vérification des tokens JWT.'
)

heading3('Ce que fait ce service :')
bullet('**Inscription** : un nouvel utilisateur crée un compte avec son nom, email et mot de passe. Le mot de passe est automatiquement chiffré (haché) avant stockage.')
bullet('**Connexion** : l\'utilisateur fournit email + mot de passe. Si corrects, le service génère un token JWT valable 24 heures.')
bullet('**Vérification de token** : les autres services appellent ce service pour vérifier si le token reçu est valide.')
bullet('**Gestion des rôles** : deux rôles existent — ADMIN (accès total) et OPERATOR (accès limité).')
bullet('**Protection brute-force** : maximum 5 tentatives de connexion par 15 minutes par adresse IP.')
bullet('**Suppression d\'utilisateur** : un ADMIN peut supprimer un compte (sauf le sien).')

heading3('Données stockées (table users) :')
add_table(
    ['Champ', 'Type', 'Description'],
    [
        ['id', 'INT', 'Identifiant unique auto-incrémenté'],
        ['name', 'VARCHAR', 'Nom complet de l\'utilisateur'],
        ['email', 'VARCHAR UNIQUE', 'Email unique, utilisé pour la connexion'],
        ['password', 'VARCHAR', 'Mot de passe haché avec bcrypt'],
        ['role', 'ENUM', 'ADMIN ou OPERATOR'],
        ['created_at', 'TIMESTAMP', 'Date de création du compte'],
    ],
    col_widths=[3, 3, 10]
)

doc.add_paragraph()

# 5.2 VEHICLE
heading2('5.2  Service Gestion des Véhicules  (port 3002)')
body(
    'Ce service gère le parc de véhicules surveillés par la plateforme. Il enregistre '
    'chaque véhicule (voiture, camion, bus...) et permet de simuler leur position GPS '
    'en temps réel. C\'est ainsi qu\'on peut visualiser les déplacements sur la carte.'
)

heading3('Ce que fait ce service :')
bullet('**Ajouter un véhicule** : enregistre la plaque, le type et le propriétaire. La plaque est automatiquement mise en majuscule et doit être unique.')
bullet('**Modifier un véhicule** : un opérateur peut modifier les informations d\'un véhicule existant.')
bullet('**Supprimer un véhicule** : réservé aux ADMIN uniquement. Supprime aussi l\'historique GPS associé.')
bullet('**Liste avec pagination** : retourne les véhicules par page (50 par défaut, max 100).')
bullet('**Enregistrer une position GPS** : stocke latitude, longitude et vitesse avec horodatage.')
bullet('**Historique des déplacements** : récupère toutes les positions enregistrées pour un véhicule.')
bullet('**Dernière position** : retourne uniquement la position la plus récente d\'un véhicule.')
bullet('**Statistiques** : nombre total de véhicules, total de positions GPS, véhicules actifs aujourd\'hui.')

heading3('Données stockées :')
add_table(
    ['Table', 'Champs principaux', 'Description'],
    [
        ['vehicles', 'id, plate, type, owner, created_at', 'Informations du véhicule'],
        ['gps_positions', 'id, vehicle_id, latitude, longitude, speed, recorded_at', 'Positions GPS enregistrées'],
    ],
    col_widths=[3, 6, 7]
)

doc.add_paragraph()

# 5.3 TRAFFIC
heading2('5.3  Service Gestion du Trafic  (port 3003)')
body(
    'Ce service est responsable de la surveillance des zones de circulation. La ville est '
    'divisée en zones géographiques définies par des coordonnées GPS. Pour chaque zone, '
    'on peut mesurer la densité du trafic, et le service classifie automatiquement le '
    'niveau de congestion.'
)

heading3('Ce que fait ce service :')
bullet('**Créer des zones** : définit une zone géographique rectangulaire avec un nom et des coordonnées GPS (lat_min, lat_max, lng_min, lng_max).')
bullet('**Enregistrer une mesure de densité** : pour une zone donnée, enregistre le nombre de véhicules détectés.')
bullet('**Classification automatique** : le système calcule seul le niveau de congestion selon la densité mesurée.')
bullet('**Zones congestionnées** : liste toutes les zones dont la dernière mesure est au niveau Élevé.')
bullet('**Résumé des zones** : retourne toutes les zones avec leur dernière mesure de densité.')
bullet('**Statistiques** : nombre total de zones, total de mesures, zones congestionnées dans la dernière heure.')

heading3('Niveaux de congestion (calculés automatiquement) :')
add_table(
    ['Niveau', 'Condition', 'Signification'],
    [
        ['Faible', 'Densité < 20 véhicules', 'Circulation fluide, aucune action requise'],
        ['Moyen', '20 ≤ Densité < 50 véhicules', 'Ralentissement modéré, surveillance recommandée'],
        ['Élevé', 'Densité ≥ 50 véhicules', 'Congestion — alerte WebSocket envoyée automatiquement'],
    ],
    col_widths=[2.5, 5, 8.5]
)

doc.add_paragraph()

# 5.4 INCIDENT
heading2('5.4  Service Gestion des Incidents  (port 3004)')
body(
    'Ce service permet de signaler et de suivre les incidents routiers. Quand un incident '
    'est déclaré, le service envoie automatiquement une notification broadcast à tous les '
    'utilisateurs de la plateforme via le service Notification, et déclenche un événement '
    'WebSocket pour les clients connectés en temps réel.'
)

heading3('Ce que fait ce service :')
bullet('**Déclarer un incident** : un opérateur signale un événement avec son type, sa description et ses coordonnées GPS.')
bullet('**Consulter les incidents** : liste tous les incidents, avec filtres optionnels par type, statut, et plage de dates.')
bullet('**Modifier le statut** : faire évoluer un incident de "Signalé" → "En cours" → "Résolu".')
bullet('**Supprimer un incident** : réservé aux ADMIN.')
bullet('**Statistiques** : nombre d\'incidents par type et par statut, total, et incidents du jour.')
bullet('**Notification automatique** : à chaque nouvel incident, tous les utilisateurs sont notifiés.')

heading3('Types d\'incidents :')
add_table(
    ['Type', 'Description'],
    [
        ['Accident', 'Collision entre véhicules ou accident de la route'],
        ['Travaux', 'Chantier de construction ou travaux de voirie'],
        ['Route fermée', 'Fermeture temporaire d\'une voie de circulation'],
        ['Embouteillage', 'Bouchon sans incident particulier, forte densité'],
    ],
    col_widths=[4, 12]
)

doc.add_paragraph()

heading3('Cycle de vie d\'un incident :')
p = doc.add_paragraph()
r = p.add_run('  [ Signalé ]  →  [ En cours ]  →  [ Résolu ]')
set_font(r, name='Courier New', size=11, bold=True, color=(30, 100, 200))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 5.5 NOTIFICATION
heading2('5.5  Service Notifications  (port 3005)')
body(
    'Ce service stocke et distribue les notifications aux utilisateurs. Il est appelé '
    'automatiquement par le Service Incidents lorsqu\'un nouvel incident est déclaré. '
    'Les notifications sont persistantes en base de données et peuvent être consultées '
    'à tout moment.'
)

heading3('Ce que fait ce service :')
bullet('**Envoi ciblé** : envoie une notification à un utilisateur spécifique (via user_id).')
bullet('**Diffusion (broadcast)** : envoie une notification à TOUS les utilisateurs enregistrés.')
bullet('**Consulter mes notifications** : retourne toutes les notifications de l\'utilisateur connecté.')
bullet('**Notifications non lues** : filtre uniquement les notifications pas encore lues.')
bullet('**Compter les non lues** : retourne le total et le nombre de non lues.')
bullet('**Marquer comme lue** : marque une notification spécifique comme lue.')
bullet('**Tout marquer comme lu** : marque toutes les notifications de l\'utilisateur comme lues.')
bullet('**Supprimer** : supprime une notification de l\'historique.')

separator()

# ════════════════════════════════════════════════════════════════
#  6. API GATEWAY GRAPHQL
# ════════════════════════════════════════════════════════════════
heading1('6. API Gateway GraphQL')

heading2('Qu\'est-ce que le Gateway ?')
body(
    'Le Gateway est le seul point d\'entrée de la plateforme. Au lieu d\'exposer 5 APIs '
    'séparées, il en propose une seule, unifiée, accessible sur http://localhost:4000/graphql. '
    'Il utilise Apollo Server v4 avec le standard GraphQL.'
)

heading2('Pourquoi GraphQL plutôt que REST ?')
add_table(
    ['REST classique', 'GraphQL (notre projet)'],
    [
        ['Plusieurs endpoints (/vehicles, /incidents...)', 'Un seul endpoint : /graphql'],
        ['Retourne toujours tous les champs', 'Le client demande exactement ce qu\'il veut'],
        ['Plusieurs requêtes pour plusieurs ressources', 'Une seule requête suffit'],
        ['Documentation séparée nécessaire', 'Schema auto-documenté (introspection)'],
    ],
    col_widths=[8, 8]
)

doc.add_paragraph()
heading2('Types d\'opérations GraphQL')
bullet('**Query** : lire des données (équivalent GET en REST) — ex : récupérer la liste des véhicules')
bullet('**Mutation** : modifier des données (équivalent POST/PUT/DELETE) — ex : ajouter un véhicule')

heading2('Liste complète des opérations disponibles')

heading3('Queries (lecture de données) :')
add_table(
    ['Opération', 'Description'],
    [
        ['me', 'Voir mon profil utilisateur'],
        ['getUsers', 'Liste de tous les utilisateurs (ADMIN)'],
        ['getVehicles', 'Liste des véhicules avec pagination'],
        ['getVehicle(id)', 'Détail d\'un véhicule par ID'],
        ['getVehicleHistory(id)', 'Historique GPS d\'un véhicule'],
        ['getLastPosition(id)', 'Dernière position GPS d\'un véhicule'],
        ['getVehicleStats', 'Statistiques véhicules (total, actifs aujourd\'hui)'],
        ['getZones', 'Liste des zones de circulation'],
        ['getZoneSummary', 'Zones avec leur dernière mesure de densité'],
        ['getCongestedZones', 'Zones à niveau Élevé uniquement'],
        ['getZoneMeasures(zone_id)', 'Historique des mesures d\'une zone'],
        ['getTrafficStats', 'Statistiques trafic globales'],
        ['getIncidents', 'Liste des incidents (filtrables)'],
        ['getIncident(id)', 'Détail d\'un incident'],
        ['getIncidentStats', 'Stats par type et statut'],
        ['getNotifications', 'Mes notifications'],
        ['getUnreadNotifications', 'Mes notifications non lues'],
        ['getNotificationCount', 'Nombre total et non lu'],
        ['healthCheck', 'État de santé des 5 services'],
        ['getPlatformStats', 'Statistiques globales de la plateforme'],
    ],
    col_widths=[6, 10]
)

doc.add_paragraph()
heading3('Mutations (modification de données) :')
add_table(
    ['Opération', 'Description'],
    [
        ['register', 'Créer un nouveau compte utilisateur'],
        ['login', 'Se connecter et obtenir un JWT'],
        ['deleteUser(id)', 'Supprimer un utilisateur (ADMIN)'],
        ['addVehicle', 'Ajouter un nouveau véhicule'],
        ['updateVehicle(id)', 'Modifier les infos d\'un véhicule'],
        ['deleteVehicle(id)', 'Supprimer un véhicule (ADMIN)'],
        ['addGpsPosition', 'Enregistrer une position GPS'],
        ['createZone', 'Créer une zone de circulation'],
        ['deleteZone(id)', 'Supprimer une zone (ADMIN)'],
        ['addTrafficMeasure', 'Mesurer la densité d\'une zone'],
        ['reportIncident', 'Déclarer un incident (déclenche notification + WebSocket)'],
        ['updateIncidentStatus(id)', 'Changer le statut d\'un incident'],
        ['deleteIncident(id)', 'Supprimer un incident (ADMIN)'],
        ['markNotificationRead(id)', 'Marquer une notification comme lue'],
        ['markAllNotificationsRead', 'Tout marquer comme lu'],
        ['deleteNotification(id)', 'Supprimer une notification'],
    ],
    col_widths=[6, 10]
)

doc.add_paragraph()
separator()

# ════════════════════════════════════════════════════════════════
#  7. WEBSOCKET
# ════════════════════════════════════════════════════════════════
heading1('7. WebSocket — Temps Réel')

heading2('Principe')
body(
    'En plus de GraphQL (qui fonctionne en requête-réponse), la plateforme intègre '
    'WebSocket via Socket.io. Cela permet au serveur d\'envoyer des messages aux clients '
    'sans qu\'ils n\'aient besoin de demander. C\'est la technologie utilisée par les '
    'applications de chat, les tableaux de bord en temps réel, ou ici les alertes trafic.'
)

heading2('Événements émis automatiquement')
add_table(
    ['Événement', 'Déclencheur', 'Données envoyées'],
    [
        ['new_incident', 'Quand un incident est déclaré', 'id, type, description, coordonnées GPS, statut, timestamp'],
        ['incident_updated', 'Quand le statut d\'un incident change', 'id de l\'incident, nouveau statut, timestamp'],
        ['traffic_alert', 'Quand une zone atteint le niveau Élevé', 'zone_id, nom de la zone, densité, niveau, timestamp'],
        ['welcome', 'À la connexion d\'un client WebSocket', 'Message de bienvenue, timestamp'],
    ],
    col_widths=[4, 5, 7]
)

doc.add_paragraph()

heading2('Comment tester le WebSocket')
body(
    'Pour tester les événements WebSocket en direct, on peut utiliser un simple fichier HTML '
    'dans le navigateur. Dès qu\'un incident est déclaré via GraphQL, le message apparaît '
    'immédiatement dans le navigateur sans avoir besoin de rafraîchir la page.'
)

separator()

# ════════════════════════════════════════════════════════════════
#  8. BASE DE DONNÉES
# ════════════════════════════════════════════════════════════════
heading1('8. Base de données')

heading2('Principe d\'isolation')
body(
    'La plateforme utilise 5 bases de données MySQL séparées, une par service. '
    'Cette séparation garantit qu\'une panne ou modification dans une base n\'affecte '
    'pas les autres services.'
)

add_table(
    ['Base de données', 'Service associé', 'Tables'],
    [
        ['db_auth', 'Auth Service', 'users'],
        ['db_vehicles', 'Vehicle Service', 'vehicles, gps_positions'],
        ['db_traffic', 'Traffic Service', 'zones, traffic_measures'],
        ['db_incidents', 'Incident Service', 'incidents'],
        ['db_notifications', 'Notification Service', 'notifications'],
    ],
    col_widths=[4, 5, 7]
)

doc.add_paragraph()

heading2('Comment créer les bases de données')
bullet('Démarrer XAMPP et activer MySQL')
bullet('Ouvrir phpMyAdmin (http://localhost/phpmyadmin)')
bullet('Cliquer sur l\'onglet "Importer"')
bullet('Sélectionner le fichier backend/database/schema.sql')
bullet('Cliquer "Exécuter" — les 5 bases sont créées automatiquement')

separator()

# ════════════════════════════════════════════════════════════════
#  9. SÉCURITÉ
# ════════════════════════════════════════════════════════════════
heading1('9. Sécurité')

add_table(
    ['Mécanisme', 'Description', 'Où appliqué'],
    [
        ['JWT (JSON Web Token)', 'Token signé envoyé à chaque requête pour prouver l\'identité', 'Tous les services'],
        ['bcrypt', 'Algorithme de hachage du mot de passe (irréversible)', 'Service Auth'],
        ['Rate Limiting', 'Maximum 5 tentatives de connexion par 15 minutes par IP', 'Route /login'],
        ['Contrôle des rôles', 'Les ADMIN ont accès aux suppressions, les OPERATOR non', 'Tous les services'],
        ['Validation des entrées', 'Vérification du format email, longueur du mot de passe, coordonnées GPS valides', 'Tous les services'],
        ['CORS', 'Autorise les appels depuis n\'importe quelle origine', 'Tous les services'],
        ['Headers de sécurité', 'X-Content-Type-Options pour éviter l\'injection de contenu', 'Tous les services'],
        ['Token expiration', 'Le JWT expire après 24 heures — reconnexion obligatoire', 'Service Auth'],
    ],
    col_widths=[4, 8, 4]
)

doc.add_paragraph()
separator()

# ════════════════════════════════════════════════════════════════
#  10. DOCKER
# ════════════════════════════════════════════════════════════════
heading1('10. Docker & Déploiement')

heading2('Qu\'est-ce que Docker ?')
body(
    'Docker permet d\'emballer une application et toutes ses dépendances dans un '
    '"conteneur" — un environnement isolé qui fonctionne de façon identique sur '
    'n\'importe quel ordinateur. Docker Compose permet de gérer plusieurs conteneurs '
    'ensemble.'
)

heading2('Notre configuration Docker')
body(
    'Le fichier docker-compose.yml lance en une seule commande l\'ensemble de la '
    'plateforme : MySQL + les 5 services + le Gateway, tous connectés via un réseau '
    'interne Docker.'
)

add_table(
    ['Conteneur', 'Image / Base', 'Port exposé'],
    [
        ['traffic_mysql', 'MySQL 8.0 officiel', '3306'],
        ['traffic_auth', 'Node.js 18 Alpine', '3001'],
        ['traffic_vehicle', 'Node.js 18 Alpine', '3002'],
        ['traffic_traffic', 'Node.js 18 Alpine', '3003'],
        ['traffic_incident', 'Node.js 18 Alpine', '3004'],
        ['traffic_notification', 'Node.js 18 Alpine', '3005'],
        ['traffic_gateway', 'Node.js 18 Alpine', '4000'],
    ],
    col_widths=[5, 5, 3.5]
)

doc.add_paragraph()

heading2('Commandes Docker')
bullet('**Lancer tout** : docker-compose up --build')
bullet('**Arrêter tout** : docker-compose down')
bullet('**Voir les logs** : docker-compose logs -f gateway')
bullet('**Voir l\'état** : docker-compose ps')

separator()

# ════════════════════════════════════════════════════════════════
#  11. GUIDE DE TEST
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1('11. Guide de test complet')

body(
    'Toutes les requêtes se font via Apollo Sandbox sur http://localhost:4000/graphql. '
    'Copier chaque bloc dans la zone de gauche et cliquer le bouton ▶ pour exécuter.'
)

heading2('Étape 1 — Inscription (aucun token requis)')
p = doc.add_paragraph()
r = p.add_run(
    'mutation {\n'
    '  register(\n'
    '    name: "Ali Ben Salah"\n'
    '    email: "ali@tekup.tn"\n'
    '    password: "pass123"\n'
    '    role: "ADMIN"\n'
    '  ) { message }\n'
    '}'
)
set_font(r, name='Courier New', size=9, color=(0, 80, 0))

heading2('Étape 2 — Connexion (copier le token reçu !)')
p = doc.add_paragraph()
r = p.add_run(
    'mutation {\n'
    '  login(email: "ali@tekup.tn", password: "pass123") {\n'
    '    token\n'
    '    user { id name role }\n'
    '  }\n'
    '}'
)
set_font(r, name='Courier New', size=9, color=(0, 80, 0))

body('➡ Dans "Headers" (en bas d\'Apollo) : ajouter   authorization: Bearer VOTRE_TOKEN')

heading2('Étape 3 — Vérification de santé des services')
p = doc.add_paragraph()
r = p.add_run(
    'query {\n'
    '  healthCheck { service status timestamp }\n'
    '  getPlatformStats {\n'
    '    total_users total_vehicles total_zones\n'
    '    total_incidents total_notifications congested_zones\n'
    '  }\n'
    '}'
)
set_font(r, name='Courier New', size=9, color=(0, 80, 0))

heading2('Étape 4 — Ajouter et consulter des véhicules')
p = doc.add_paragraph()
r = p.add_run(
    '# Ajouter\n'
    'mutation { addVehicle(plate: "TN-1234-A", type: "Voiture", owner: "Ali") { message } }\n\n'
    '# Consulter\n'
    'query { getVehicles { data { id plate type owner } total } }\n\n'
    '# Position GPS\n'
    'mutation {\n'
    '  addGpsPosition(vehicle_id: 1, latitude: 36.8189, longitude: 10.1658, speed: 60.5)\n'
    '  { message }\n'
    '}\n\n'
    '# Statistiques véhicules\n'
    'query { getVehicleStats { total_vehicles total_positions active_today } }'
)
set_font(r, name='Courier New', size=9, color=(0, 80, 0))

heading2('Étape 5 — Zones et trafic')
p = doc.add_paragraph()
r = p.add_run(
    '# Créer une zone\n'
    'mutation {\n'
    '  createZone(name: "Centre-Ville", lat_min: 36.80, lat_max: 36.85,\n'
    '    lng_min: 10.15, lng_max: 10.20) { message }\n'
    '}\n\n'
    '# Mesure de densité (≥50 = Élevé → alerte WebSocket)\n'
    'mutation { addTrafficMeasure(zone_id: 1, density: 75) { message } }\n\n'
    '# Résumé et zones congestionnées\n'
    'query {\n'
    '  getZoneSummary { id name density level measured_at }\n'
    '  getCongestedZones { zone_id name density level }\n'
    '}'
)
set_font(r, name='Courier New', size=9, color=(0, 80, 0))

heading2('Étape 6 — Incidents (déclenche WebSocket + notification auto)')
p = doc.add_paragraph()
r = p.add_run(
    '# Déclarer un incident\n'
    'mutation {\n'
    '  reportIncident(\n'
    '    type: "Accident"\n'
    '    description: "Collision sur Avenue Bourguiba"\n'
    '    latitude: 36.8190\n'
    '    longitude: 10.1660\n'
    '  ) { message }\n'
    '}\n\n'
    '# Voir les incidents\n'
    'query { getIncidents { id type description status created_at } }\n\n'
    '# Stats par type/statut\n'
    'query { getIncidentStats { total today by_type { type count } by_status { status count } } }\n\n'
    '# Changer le statut\n'
    'mutation { updateIncidentStatus(id: 1, status: "En cours") { message } }'
)
set_font(r, name='Courier New', size=9, color=(0, 80, 0))

heading2('Étape 7 — Notifications')
p = doc.add_paragraph()
r = p.add_run(
    '# Voir mes notifications (créées auto par l\'incident)\n'
    'query { getNotifications { id message is_read created_at } }\n\n'
    '# Compter les non lues\n'
    'query { getNotificationCount { total unread } }\n\n'
    '# Tout marquer comme lu\n'
    'mutation { markAllNotificationsRead { message } }'
)
set_font(r, name='Courier New', size=9, color=(0, 80, 0))

separator()

# ════════════════════════════════════════════════════════════════
#  12. RÉSUMÉ
# ════════════════════════════════════════════════════════════════
heading1('12. Résumé des fonctionnalités')

add_table(
    ['Critère du prof', 'Implémenté', 'Détails'],
    [
        ['Architecture microservices', '✅ OUI', '5 services indépendants + Gateway'],
        ['GraphQL obligatoire', '✅ OUI', 'Apollo Server v4, 20 queries + 16 mutations'],
        ['Base de données relationnelle', '✅ OUI', 'MySQL 8, 5 bases isolées'],
        ['Validation des données', '✅ OUI', 'Format email, longueur mdp, coordonnées GPS, énumérations'],
        ['Gestion des erreurs', '✅ OUI', 'Codes HTTP corrects, messages clairs'],
        ['JWT Authentication', '✅ OUI', 'Token 24h, vérification inter-services'],
        ['Docker Compose (bonus)', '✅ OUI', '7 conteneurs, healthcheck MySQL'],
        ['WebSocket temps réel (bonus)', '✅ OUI', 'Socket.io, 3 événements automatiques'],
        ['Frontend React (bonus)', '✅ OUI', 'Vite + React, dashboard complet'],
        ['Historique GitHub', '✅ OUI', 'Commit initial complet avec description'],
        ['Collection Postman', '✅ OUI', 'Fichier JSON importable'],
        ['Diagrammes UML', '✅ OUI', 'Fichier .puml dans /docs'],
    ],
    col_widths=[6, 2.5, 7.5]
)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— Fin du document —')
set_font(r, size=10, color=(150, 150, 150))

# ════════════════════════════════════════════════════════════════
#  SAVE
# ════════════════════════════════════════════════════════════════
output = r'D:\TEK-UP\2 annee\semester 2\Service Web\Projet\traffic-platform\Traffic_Platform_Documentation.docx'
doc.save(output)
print(f'[OK] Document genere : {output}')
